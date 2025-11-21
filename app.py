import streamlit as st
import os, json, re, requests, time, tempfile
from datetime import datetime, timedelta, date
from lxml import etree
from io import BytesIO
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from msal import ConfidentialClientApplication
from openai import AzureOpenAI
import PyPDF2
import docx
import pandas as pd
import base64
from PIL import Image

# ------------------- TRANSLATIONS -------------------
TRANSLATIONS = {
    "en": {
        # Top bar
        "title": "TED Scraper & AI Assistant",
        "language": "Language",
        
        # Tab names
        "tab_scraper": "📄 TED Scraper",
        "tab_assistant": "💬 AI Assistant",
        
        # Scraper section
        "scraper_header": "📄 TED EU Notice Scraper",
        "scraper_subtitle": "Search and filter TED procurement notices before downloading.",
        "instructions_header": "ℹ️ How this works / Instructions",
        "features_title": "**FEATURES:**",
        "feature_1": "- 🔍 Search by **keywords** (single or multi-word) OR **CPV codes** OR **both**",
        "feature_2": "- 👀 **Preview results** before downloading",
        "feature_3": "- 🎯 **Multi-select filters** - Select multiple Beschaffer or Regions with checkboxes!",
        "feature_4": "- ⬇️ **Download only filtered data**",
        "keywords_examples": "**Keywords Examples:**",
        "example_1": "- Single word: `construction` ✅",
        "example_2": "- Multi-word: `project management` ✅",
        "multiselect_info": "**Multi-Select Filters:**",
        "multiselect_1": "- Click dropdown and tick multiple options",
        "multiselect_2": "- Filter by multiple contractors or locations at once",
        
        # Search inputs
        "search_criteria": "🔍 Search Criteria",
        "keywords_label": "🔤 Keywords (single or multi-word)",
        "keywords_placeholder": "e.g., project management, quality assurance",
        "keywords_help": "Single words or phrases work!",
        "cpv_label": "🏷️ CPV Codes (space separated)",
        "cpv_help": "Classification codes. Leave empty to search by keywords only.",
        "country_label": "🌍 Buyer Country (ISO Alpha-3)",
        "country_help": "e.g., DEU, FRA, ITA, ESP",
        "date_start_label": "📆 Publication Start",
        "date_end_label": "📆 Publication End",
        "search_button": "🔍 Search Notices",
        
        # Errors and warnings
        "error_no_keywords": "❌ Please enter either keywords or CPV codes (or both)!",
        "searching": "Searching TED database... This may take a few minutes.",
        "success_found": "✅ Found {count} notices!",
        "warning_no_results": "⚠️ No results found. Try adjusting your search criteria.",
        "error_search": "❌ Error during search: {error}",
        
        # Results section
        "results_header": "📊 Search Results",
        "total_results": "📈 Total Results: **{count}** notices",
        "filter_results": "🎯 Filter Results",
        "filter_beschaffer": "✅ Filter by Beschaffer (Multi-select)",
        "filter_beschaffer_help": "Select multiple contractors using checkboxes",
        "filter_region": "✅ Filter by Region (Multi-select)",
        "filter_region_help": "Select multiple locations using checkboxes",
        "filter_volume": "Min Volume (EUR)",
        "filter_volume_placeholder": "e.g., 100000",
        "filter_projektstart": "🗓️ Projectstart",
        "filter_projektstart_help": "Filter notices with project start date on or after this date",
        "filter_projektende": "🗓️ Projectend",
        "filter_projektende_help": "Filter notices with project end date on or before this date",
        "filter_frist": "⏰ Frist Abgabedatum",
        "filter_frist_help": "Filter notices with submission deadline on or after this date",
        "filtered_results": "🎯 Filtered Results: **{count}** notices",
        "warning_volume": "⚠️ Invalid volume filter",
        
        # Download buttons
        "download_filtered": "⬇️ Download Filtered Results ({count} notices)",
        "download_all": "⬇️ Download All Results ({count} notices)",
        
        # Chatbot section
        "config_header": "## 🔑 Configuration",
        "azure_connected": "✅ Azure AI Connected",
        "azure_warning": "⚠️ Azure credentials missing",
        "doc_library": "## 📚 Document Library",
        "doc_optional": "Optional: Upload files for context",
        "clear_chat": "🗑️ Clear Chat",
        "azure_error": "❌ Azure AI Foundry credentials not configured!",
        "azure_info": "**Add to `.streamlit/secrets.toml`:**",
        "welcome_header": "👋 **Welcome to JKM AI Assistant!**",
        "welcome_text": "I am your AI assistant and can help you with various tasks.",
        "possibilities": "**Possibilities:**",
        "possibility_1": "- 💬 Answer general questions",
        "possibility_2": "- 📄 Analyze documents (PDF, Word, TXT)",
        "possibility_3": "- 🔍 Review tenders",
        "possibility_4": "- ✍️ Write and translate texts",
        "ask_question": "Just ask me a question!",
        "file_upload": "📎 Drag and drop file here or click to browse",
        "file_help": "Upload documents, Excel files, or images",
        "file_added": "✅ {filename} added",
        "chat_input": "Message JKM AI Assistant...",
        "thinking": "💭 AI is thinking",
        "error_check_config": "Please check your Azure configuration in secrets.toml",
        "processing": "Processing {filename}...",
        "query_label": "🔍 Query: `{query}`",
    },
    "de": {
        # Top bar
        "title": "TED Scraper & AI Assistent",
        "language": "Sprache",
        
        # Tab names
        "tab_scraper": "📄 TED Scraper",
        "tab_assistant": "💬 KI-Assistent",
        
        # Scraper section
        "scraper_header": "📄 TED EU Ausschreibungs-Scraper",
        "scraper_subtitle": "Durchsuchen und filtern Sie TED-Ausschreibungen vor dem Herunterladen.",
        "instructions_header": "ℹ️ So funktioniert es / Anleitung",
        "features_title": "**FUNKTIONEN:**",
        "feature_1": "- 🔍 Suche nach **Schlüsselwörtern** (einzeln oder mehrere Wörter) ODER **CPV-Codes** ODER **beides**",
        "feature_2": "- 👀 **Vorschau der Ergebnisse** vor dem Herunterladen",
        "feature_3": "- 🎯 **Multi-Select-Filter** - Wählen Sie mehrere Beschaffer oder Regionen mit Checkboxen!",
        "feature_4": "- ⬇️ **Nur gefilterte Daten herunterladen**",
        "keywords_examples": "**Schlüsselwort-Beispiele:**",
        "example_1": "- Einzelwort: `construction` ✅",
        "example_2": "- Mehrere Wörter: `project management` ✅",
        "multiselect_info": "**Multi-Select-Filter:**",
        "multiselect_1": "- Dropdown anklicken und mehrere Optionen auswählen",
        "multiselect_2": "- Nach mehreren Auftraggebern oder Standorten gleichzeitig filtern",
        
        # Search inputs
        "search_criteria": "🔍 Suchkriterien",
        "keywords_label": "🔤 Schlüsselwörter (einzeln oder mehrere)",
        "keywords_placeholder": "z.B., Projektmanagement, Qualitätssicherung",
        "keywords_help": "Einzelne Wörter oder Phrasen funktionieren!",
        "cpv_label": "🏷️ CPV-Codes (durch Leerzeichen getrennt)",
        "cpv_help": "Klassifikationscodes. Leer lassen, um nur nach Schlüsselwörtern zu suchen.",
        "country_label": "🌍 Auftraggeber-Land (ISO Alpha-3)",
        "country_help": "z.B., DEU, FRA, ITA, ESP",
        "date_start_label": "📆 Veröffentlichung Start",
        "date_end_label": "📆 Veröffentlichung Ende",
        "search_button": "🔍 Ausschreibungen suchen",
        
        # Errors and warnings
        "error_no_keywords": "❌ Bitte geben Sie entweder Schlüsselwörter oder CPV-Codes ein (oder beides)!",
        "searching": "Durchsuche TED-Datenbank... Dies kann einige Minuten dauern.",
        "success_found": "✅ {count} Ausschreibungen gefunden!",
        "warning_no_results": "⚠️ Keine Ergebnisse gefunden. Versuchen Sie, Ihre Suchkriterien anzupassen.",
        "error_search": "❌ Fehler bei der Suche: {error}",
        
        # Results section
        "results_header": "📊 Suchergebnisse",
        "total_results": "📈 Gesamtergebnisse: **{count}** Ausschreibungen",
        "filter_results": "🎯 Ergebnisse filtern",
        "filter_beschaffer": "✅ Nach Beschaffer filtern (Mehrfachauswahl)",
        "filter_beschaffer_help": "Wählen Sie mehrere Auftraggeber mit Checkboxen aus",
        "filter_region": "✅ Nach Region filtern (Mehrfachauswahl)",
        "filter_region_help": "Wählen Sie mehrere Standorte mit Checkboxen aus",
        "filter_volume": "Min. Volumen (EUR)",
        "filter_volume_placeholder": "z.B., 100000",
        "filter_projektstart": "🗓️ Projektstart",
        "filter_projektstart_help": "Ausschreibungen mit Projektstart an oder nach diesem Datum filtern",
        "filter_projektende": "🗓️ Projektende",
        "filter_projektende_help": "Ausschreibungen mit Projektende an oder vor diesem Datum filtern",
        "filter_frist": "⏰ Abgabefrist",
        "filter_frist_help": "Ausschreibungen mit Abgabefrist an oder nach diesem Datum filtern",
        "filtered_results": "🎯 Gefilterte Ergebnisse: **{count}** Ausschreibungen",
        "warning_volume": "⚠️ Ungültiger Volumenfilter",
        
        # Download buttons
        "download_filtered": "⬇️ Gefilterte Ergebnisse herunterladen ({count} Ausschreibungen)",
        "download_all": "⬇️ Alle Ergebnisse herunterladen ({count} Ausschreibungen)",
        
        # Chatbot section
        "config_header": "## 🔑 Konfiguration",
        "azure_connected": "✅ Azure AI Verbunden",
        "azure_warning": "⚠️ Azure-Anmeldedaten fehlen",
        "doc_library": "## 📚 Dokumentenbibliothek",
        "doc_optional": "Optional: Dateien für Kontext hochladen",
        "clear_chat": "🗑️ Chat leeren",
        "azure_error": "❌ Azure AI Foundry-Anmeldedaten nicht konfiguriert!",
        "azure_info": "**Zu `.streamlit/secrets.toml` hinzufügen:**",
        "welcome_header": "👋 **Willkommen beim JKM AI Assistent!**",
        "welcome_text": "Ich bin Ihr KI-Assistent und kann Ihnen bei verschiedenen Aufgaben helfen.",
        "possibilities": "**Möglichkeiten:**",
        "possibility_1": "- 💬 Allgemeine Fragen beantworten",
        "possibility_2": "- 📄 Dokumente analysieren (PDF, Word, TXT)",
        "possibility_3": "- 🔍 Ausschreibungen prüfen",
        "possibility_4": "- ✍️ Texte schreiben und übersetzen",
        "ask_question": "Stellen Sie mir einfach eine Frage!",
        "file_upload": "📎 Datei hier ablegen oder zum Durchsuchen klicken",
        "file_help": "Dokumente, Excel-Dateien oder Bilder hochladen",
        "file_added": "✅ {filename} hinzugefügt",
        "chat_input": "Nachricht an JKM AI Assistent...",
        "thinking": "💭 KI denkt nach",
        "error_check_config": "Bitte überprüfen Sie Ihre Azure-Konfiguration in secrets.toml",
        "processing": "Verarbeite {filename}...",
        "query_label": "🔍 Abfrage: `{query}`",
    }
}

def t(key, **kwargs):
    """Translation helper function"""
    lang = st.session_state.get("language", "en")
    text = TRANSLATIONS.get(lang, TRANSLATIONS["en"]).get(key, key)
    if kwargs:
        text = text.format(**kwargs)
    return text

# ------------------- CONFIGURATION -------------------
def get_secret(key, default=""):
    """Safely get secrets with fallback"""
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default

CLIENT_ID = get_secret("CLIENT_ID")
CLIENT_SECRET = get_secret("CLIENT_SECRET")
TENANT_ID = get_secret("TENANT_ID")
REDIRECT_URI = get_secret("REDIRECT_URI", "http://localhost:8501")

AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}" if TENANT_ID else ""
SCOPE = ["https://graph.microsoft.com/User.Read"]

API = "https://api.ted.europa.eu/v3/notices/search"

# Avatars
JKM_LOGO_URL = "https://www.xing.com/imagecache/public/scaled_original_image/eyJ1dWlkIjoiMGE2MTk2MTYtODI4Zi00MWZlLWEzN2ItMjczZGM2ODc5MGJmIiwiYXBwX2NvbnRleHQiOiJlbnRpdHktcGFnZXMiLCJtYXhfd2lkdGgiOjMyMCwibWF4X2hlaWdodCI6MzIwfQ?signature=a21e5c1393125a94fc9765898c25d73a064665dc3aacf872667c902d7ed9c3f9"
BOT_AVATAR_URL = "https://raw.githubusercontent.com/PratikSondkarJKM/AkquiseWescraper/refs/heads/main/botavatar.svg"

# ------------------- AUTHENTICATION -------------------
def build_msal_app():
    if not CLIENT_ID or not CLIENT_SECRET or not TENANT_ID:
        st.error("❌ Microsoft OAuth credentials not configured!")
        st.info("""
        Please create `.streamlit/secrets.toml` in your project directory with:
        
        ```
        CLIENT_ID = "your-client-id"
        CLIENT_SECRET = "your-client-secret"
        TENANT_ID = "your-tenant-id"
        REDIRECT_URI = "http://localhost:8501"
        ```
        """)
        st.stop()
    
    return ConfidentialClientApplication(
        client_id=CLIENT_ID,
        authority=AUTHORITY,
        client_credential=CLIENT_SECRET
    )

def fetch_token(auth_code):
    msal_app = build_msal_app()
    return msal_app.acquire_token_by_authorization_code(auth_code, scopes=SCOPE, redirect_uri=REDIRECT_URI)

def login_button():
    msal_app = build_msal_app()
    auth_url = msal_app.get_authorization_request_url(SCOPE, redirect_uri=REDIRECT_URI)
    st.markdown("""
    <style>
    .block-container { padding: 0 !important; max-width: 100vw !important; }
    .center-root {
        min-height: 100vh; width: 100vw;
        display: flex; flex-direction: column; align-items: center; justify-content: center;
        background: linear-gradient(120deg, #eaf6fb 0%, #f3e9f5 100%);
    }
    .jkm-logo {
        height: 72px; margin-bottom: 14px; border-radius: 14px; box-shadow: 0 2px 14px rgba(70,80,120,0.08);
        background: #fff; display: block;
    }
    .app-title {
        font-family: 'Segoe UI', Arial,sans-serif;
        font-size: 2.3em; text-align: center; font-weight: 800; color: #283044; margin-bottom: 10px; margin-top: 4px;
    }
    .welcome-text {
        font-size: 1.07em; color: #505A69; margin-bottom: 22px; margin-top: 0; text-align: center;
    }
    .login-card {
        width: 375px; padding: 38px 34px 31px 34px; background: #fff; border-radius: 18px;
        box-shadow: 0 8px 32px rgba(50,72,140,.13); text-align: center; margin-top: 6px;
    }
    .microsoft-logo {
        height: 44px; margin-bottom: 16px; display:block; margin-left:auto; margin-right:auto;
    }
    .login-button {
        display: block; width: 100%; padding: 17px 0 14px 0; margin: 28px 0 18px 0; font-size: 17px;
        background-color: #0078d7; color: #fff !important; border: none; border-radius: 7px;
        cursor: pointer; text-decoration: none; font-weight: 600; transition: background 0.18s; outline: none;
    }
    .login-button:hover {
        background-color: #005fa1; color: #fff !important; text-decoration: none;
    }
    </style>
    """, unsafe_allow_html=True)
    st.markdown(f"""
    <div class="center-root">
        <img src="{JKM_LOGO_URL}" class="jkm-logo" alt="JKM Consult Logo"/>
        <div class="app-title">TED Scraper & AI Assistant</div>
        <div class="welcome-text">
            Welcome! Access project info securely.<br>
            Login with Microsoft to continue.
        </div>
        <div class="login-card">
            <img src="https://upload.wikimedia.org/wikipedia/commons/4/44/Microsoft_logo.svg" class="microsoft-logo" alt="Microsoft Logo"/>
            <h2 style="margin-bottom: 9px; font-size: 1.26em;">Sign in</h2>
            <p style="font-size: 1em; color: #232b39; margin-bottom: 9px;">
                to continue to <b>TED Scraper & AI Assistant</b>
            </p>
            <a href="{auth_url}" class="login-button">
                Sign in with Microsoft
            </a>
            <p style="margin-top: 32px; font-size: 0.98em; color: #888;">
                Your credentials are always handled by Microsoft.<br>
                We never see or store your password.
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

def auth_flow():
    params = st.query_params
    if "code" in params and "user_token" not in st.session_state:
        code = params["code"]
        if isinstance(code, list):
            code = code[0]
        token_data = fetch_token(code)
        if "access_token" in token_data:
            st.session_state["user_token"] = token_data["access_token"]
            st.query_params.clear()
            st.rerun()
        else:
            st.error("Microsoft login failed. Please try again.")
            st.stop()
    if "user_token" not in st.session_state:
        login_button()
        st.stop()
    return True

# ---------------- TED SCRAPER FUNCTIONS ----------------
def fetch_all_notices_to_json(cpv_codes, keywords, date_start, date_end, buyer_country, json_file):
    """Fetch TED notices with CORRECT TED API v3 query syntax"""
    query_parts = []
    
    query_parts.append(f"(publication-date >={date_start}<={date_end})")
    query_parts.append(f"(buyer-country IN ({buyer_country}))")
    
    if cpv_codes and cpv_codes.strip():
        query_parts.append(f"(classification-cpv IN ({cpv_codes}))")
    
    if keywords and keywords.strip():
        clean_keywords = keywords.strip().replace('"', '').replace("'", "")
        query_parts.append(f"(FT~({clean_keywords}))")
    
    query_parts.append("(notice-type IN (pin-cfc-standard pin-cfc-social qu-sy cn-standard cn-social subco cn-desg))")
    
    query = " AND ".join(query_parts)
    st.info(t("query_label", query=query))
    
    payload = {
        "query": query,
        "fields": ["publication-number", "links"],
        "scope": "ACTIVE",
        "checkQuerySyntax": False,
        "paginationMode": "PAGE_NUMBER",
        "page": 1,
        "limit": 100
    }
    
    s = requests.Session()
    s.headers.update({
        "Accept": "application/json",
        "Content-Type": "application/json"
    })
    
    all_notices = []
    page = 1
    
    while True:
        body = dict(payload)
        body["page"] = page
        
        try:
            r = s.post(API, json=body, timeout=60)
            
            if r.status_code != 200:
                st.error(f"❌ API Error {r.status_code}")
                st.code(r.text[:500])
                r.raise_for_status()
            
            data = r.json()
            notices = data.get("results") or data.get("items") or data.get("notices") or []
            
            if not notices:
                break
            
            all_notices.extend(notices)
            
            total = data.get("total") or data.get("totalCount")
            if not notices or (total and page * payload["limit"] >= total):
                break
            
            page += 1
            time.sleep(0.3)
            
        except requests.exceptions.HTTPError as e:
            st.error(f"❌ HTTP Error: {e}")
            st.code(f"Query: {query}")
            raise
        except Exception as e:
            st.error(f"❌ Unexpected error: {e}")
            raise
    
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump({"notices": all_notices}, f, ensure_ascii=False, indent=2)
    
    return len(all_notices)

def _get_links_block(notice: dict) -> dict:
    links = notice.get("links") or {}
    if isinstance(links, dict) and "links" in links and isinstance(links["links"], dict):
        links = links["links"]
    if isinstance(links, dict):
        return { (k.lower() if isinstance(k,str) else k): v for k,v in links.items() }
    return {}

def _extract_xml_urls_from_notice(notice: dict) -> list:
    block = _get_links_block(notice)
    xml_block = block.get("xml")
    urls = []
    if isinstance(xml_block, dict):
        for k,v in xml_block.items():
            if isinstance(k,str) and k.lower()=="mul" and v:
                urls.append(v)
        for k,v in xml_block.items():
            if isinstance(k,str) and k.lower()!="mul" and v:
                urls.append(v)
    elif isinstance(xml_block, str) and xml_block:
        urls.append(xml_block)
    return urls

def fetch_notice_xml(session: requests.Session, pubno: str, notice: dict) -> bytes:
    xml_headers = {"Accept":"application/xml","User-Agent":"Mozilla/5.0"}
    for url in _extract_xml_urls_from_notice(notice):
        try:
            r = session.get(url, headers=xml_headers, timeout=60)
            if r.status_code == 200 and r.content.strip():
                return r.content
        except requests.RequestException:
            pass
    for lang in ("en","de","fr"):
        url = f"https://ted.europa.eu/{lang}/notice/{pubno}/xml"
        try:
            r = session.get(url, headers=xml_headers, timeout=60)
            if r.status_code == 200 and r.content.strip():
                return r.content
        except requests.RequestException:
            pass
    detail_url = f"https://ted.europa.eu/en/notice/-/detail/{pubno}"
    try:
        html = session.get(detail_url, headers={"User-Agent":"Mozilla/5.0"}, timeout=60).text
        m = re.search(r'https://ted\.europa\.eu/(?:en|de|fr)/notice/' + re.escape(pubno) + r'/xml', html)
        if m:
            r = session.get(m.group(0), headers=xml_headers, timeout=60)
            if r.status_code == 200 and r.content.strip():
                return r.content
    except requests.RequestException:
        pass
    raise RuntimeError(f"No XML found for {pubno}")

def _first_text(nodes):
    for n in nodes or []:
        t = (n.text or "").strip()
        if t:
            return t
    return ""

def _norm_date(d: str) -> str:
    if not d:
        return ""
    d = d.rstrip("Zz")
    return d.split("T")[0].split("+")[0]

def _clean_title(raw: str) -> str:
    if not raw: return ""
    return re.sub(r"^\s*\d{4}[-_]\d{5,}[\s_\-–:]+", "", raw.strip())

def _parse_iso_date(d: str):
    try:
        return datetime.strptime(d, "%Y-%m-%d")
    except Exception:
        return None

def _duration_to_days(val: str, unit: str) -> int or None:
    if not val:
        return None
    try:
        num = float(str(val).strip().replace(",", "."))
    except Exception:
        return None
    u = (unit or "").upper()
    if u in ("DAY","D","DAYS"):
        return int(round(num))
    if u in ("MON","M","MONTH","MONTHS"):
        return int(round(num * 30))
    if u in ("ANN","Y","YEAR","YEARS"):
        return int(round(num * 365))
    return None

def parse_xml_fields(xml_bytes: bytes) -> dict:
    parser = etree.XMLParser(recover=True, huge_tree=True)
    root = etree.parse(BytesIO(xml_bytes), parser)
    ns = {k: v for k, v in (root.getroot().nsmap or {}).items() if k}
    ns.setdefault("cbc","urn:oasis:names:specification:ubl:schema:xsd:CommonBasicComponents-2")
    ns.setdefault("cac","urn:oasis:names:specification:ubl:schema:xsd:CommonAggregateComponents-2")
    ns.setdefault("efac","http://data.europa.eu/p27/eforms-ubl-extension-aggregate-components/1")
    ns.setdefault("efbc","http://data.europa.eu/p27/eforms-ubl-extension-basic-components/1")

    out = {}
    out["Beschaffer"] = _first_text(
        root.xpath(".//cac:ContractingParty//cac:PartyName/cbc:Name", namespaces=ns)
        or root.xpath(".//efac:Organizations//efac:Company/cac:PartyName/cbc:Name", namespaces=ns)
    )
    out["Projektbezeichnung"] = _clean_title(
        _first_text(root.xpath(".//cac:ProcurementProject/cbc:Name | .//cbc:Title | .//efbc:Title", namespaces=ns))
    )
    out["Ort/Region"] = _first_text(root.xpath("//cac:PostalAddress[1]/cbc:CityName", namespaces=ns))
    out["Vergabeplattform"] = _first_text(
        root.xpath(".//cbc:AccessToolsURI | .//cbc:WebsiteURI | .//cbc:URI | .//cbc:EndpointID", namespaces=ns)
    )
    pub_id = _first_text(root.xpath(".//efbc:NoticePublicationID[@schemeName='ojs-notice-id']", namespaces=ns))
    out["Ted-Link"] = f"https://ted.europa.eu/en/notice/-/detail/{pub_id}" if pub_id else ""

    start_nodes = root.xpath(
        ".//cac:ProcurementProject/cac:PlannedPeriod/cbc:StartDate "
        "| .//cac:ProcurementProjectLot//cac:ProcurementProject/cac:PlannedPeriod/cbc:StartDate",
        namespaces=ns
    )
    end_nodes = root.xpath(
        ".//cac:ProcurementProject/cac:PlannedPeriod/cbc:EndDate "
        "| .//cac:ProcurementProjectLot//cac:ProcurementProject/cac:PlannedPeriod/cbc:EndDate",
        namespaces=ns
    )
    start_norm = _norm_date(_first_text(start_nodes))
    end_norm = _norm_date(_first_text(end_nodes))

    if not start_norm and end_norm:
        dur_nodes = root.xpath(
            ".//cac:ProcurementProject/cac:PlannedPeriod/cbc:DurationMeasure "
            "| .//cac:ProcurementProjectLot//cac:ProcurementProject/cac:PlannedPeriod/cbc:DurationMeasure",
            namespaces=ns
        )
        dur_val, dur_unit = None, None
        for dn in dur_nodes:
            text_val = (dn.text or "").strip()
            unit = (dn.get("unitCode") or "").strip()
            if text_val:
                dur_val, dur_unit = text_val, unit
                break
        days = _duration_to_days(dur_val, dur_unit) if dur_val else None
        if days:
            end_dt = _parse_iso_date(end_norm)
            if end_dt:
                start_norm = (end_dt - timedelta(days=days)).strftime("%Y-%m-%d")

    out["Projektstart"] = start_norm
    out["Projektende"] = end_norm

    crit_nodes = root.xpath(
        ".//*[contains(local-name(),'SelectionCriteria') or contains(local-name(),'SelectionCriterion')]/cbc:Description",
        namespaces=ns
    )
    crit_text = " ".join((n.text or "").strip() for n in crit_nodes if (n.text or "").strip())
    crit_text = re.sub(r"\bslc-[a-z0-9\-]+\b", "", crit_text, flags=re.I).strip()
    out["Geforderte Unternehmensreferenzen"] = crit_text
    out["Geforderte Kriterien CVs"] = "CV" if re.search(
        r"\b(CV|Lebenslauf|Schlüsselpersonal|key staff|personaleinsatz)\b", crit_text, re.I
    ) else ""

    amount_nodes = root.xpath(
        ".//cbc:EstimatedOverallContractAmount | .//cbc:EstimatedOverallContractAmount/cbc:Value | .//efbc:EstimatedValue | .//cbc:PayableAmount",
        namespaces=ns
    )
    value_text = ""
    if amount_nodes:
        for node in amount_nodes:
            if node.text and node.text.strip():
                value_text = node.text.strip()
                parent = node.getparent()
                currency = node.get("currencyID") or (parent.get("currencyID") if parent is not None else None)
                if currency:
                    value_text += f" {currency}"
                break
    out["Projektvolumen"] = value_text or ""

    tender_deadline_date = _norm_date(
        _first_text(root.xpath(".//cac:TenderSubmissionDeadlinePeriod/cbc:EndDate", namespaces=ns))
    )
    if not tender_deadline_date:
        tender_deadline_date = _norm_date(
            _first_text(root.xpath(".//cac:TenderingTerms/cbc:SubmissionDeadlineDate", namespaces=ns))
        )
    if not tender_deadline_date:
        tender_deadline_date = _norm_date(
            _first_text(root.xpath(".//cac:InterestExpressionReceptionPeriod/cbc:EndDate", namespaces=ns))
        )
    if not tender_deadline_date:
        tender_deadline_date = _norm_date(
            _first_text(root.xpath(".//efac:InterestExpressionReceptionPeriod/cbc:EndDate", namespaces=ns))
        )
    participation_deadline_date = _norm_date(
        _first_text(root.xpath(".//cac:ParticipationRequestReceptionPeriod/cbc:EndDate", namespaces=ns))
    )
    if not participation_deadline_date:
        participation_deadline_date = _norm_date(
            _first_text(root.xpath(".//efac:ParticipationRequestReceptionPeriod/cbc:EndDate", namespaces=ns))
        )
    out["Frist Abgabedatum"] = tender_deadline_date or participation_deadline_date

    pub_date = _first_text(root.xpath(".//efbc:PublicationDate", namespaces=ns))
    if not pub_date:
        pub_date = _first_text(root.xpath(".//cbc:PublicationDate", namespaces=ns))
    out["Veröffentlichung Datum"] = _norm_date(pub_date)

    cpv_codes_set = set()
    main_cpv_nodes = root.xpath(".//cac:MainCommodityClassification/cbc:ItemClassificationCode", namespaces=ns)
    for node in main_cpv_nodes:
        if node.text:
            cpv_codes_set.add(node.text.strip())
    add_cpv_nodes = root.xpath(".//cac:AdditionalCommodityClassification/cbc:ItemClassificationCode", namespaces=ns)
    for node in add_cpv_nodes:
        if node.text:
            cpv_codes_set.add(node.text.strip())
    out["CPV Codes"] = ", ".join(sorted(cpv_codes_set))

    lots = root.xpath(".//cac:ProcurementProjectLot", namespaces=ns)
    lot_names = []
    for lot in lots:
        lot_name = lot.xpath(".//cac:ProcurementProject/cbc:Name", namespaces=ns)
        if lot_name and len(lot_name) > 0:
            text = lot_name[0].text.strip()
            if text:
                lot_names.append(text)
    out["Leistungen/Rollen"] = "; ".join(lot_names)

    return out

def main_scraper(cpv_codes, keywords, date_start, date_end, buyer_country):
    """Modified to return rows instead of saving to Excel directly"""
    temp_json = tempfile.mktemp(suffix=".json")
    
    count = fetch_all_notices_to_json(cpv_codes, keywords, date_start, date_end, buyer_country, temp_json)
    
    if count == 0:
        st.warning(t("warning_no_results"))
        return []
    
    with open(temp_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    notices = data.get("notices", [])
    
    s = requests.Session()
    rows = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, n in enumerate(notices):
        pubno = n.get("publication-number")
        if not pubno:
            continue
        
        status_text.text(f"Processing {idx+1}/{len(notices)}: {pubno}")
        progress_bar.progress((idx + 1) / len(notices))
        
        try:
            xml_bytes = fetch_notice_xml(s, pubno, n)
            fields = parse_xml_fields(xml_bytes)
            fields["publication-number"] = pubno
            fields.setdefault("Ted-Link", f"https://ted.europa.eu/en/notice/-/detail/{pubno}")
            rows.append(fields)
        except Exception as e:
            st.warning(f"⚠️ Error processing {pubno}: {e}")
        
        time.sleep(0.25)
    
    progress_bar.empty()
    status_text.empty()
    
    os.remove(temp_json)
    return rows

def save_to_excel(rows, output_excel):
    """Save filtered rows to Excel with table formatting"""
    wb = openpyxl.Workbook()
    ws = wb.active
    headers = [
        "publication-number","Beschaffer","Projektbezeichnung","Ort/Region",
        "Vergabeplattform","Ted-Link","Projektstart","Projektende",
        "Geforderte Unternehmensreferenzen","Geforderte Kriterien CVs",
        "Projektvolumen", "Frist Abgabedatum", "Veröffentlichung Datum", "CPV Codes", "Leistungen/Rollen"
    ]
    ws.append(headers)
    for r in rows:
        ws.append([r.get(h, "") for h in headers])
    last_row = len(rows) + 1
    last_col = len(headers)
    table_range = f"A1:{get_column_letter(last_col)}{last_row}"
    table = Table(displayName="Teddata", ref=table_range)
    style = TableStyleInfo(
        name="TableStyleMedium9",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    table.tableStyleInfo = style
    ws.add_table(table)
    wb.save(output_excel)

# ---------------- CHATBOT FUNCTIONS (keep all unchanged) ----------------
def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_from_txt(file):
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        return f"Error reading TXT: {str(e)}"

def extract_text_from_excel(file):
    try:
        file_extension = file.name.split('.')[-1].lower()
        if file_extension == 'csv':
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)
        text = f"Excel File: {file.name}\n"
        text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
        text += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
        text += "Data Preview (first 50 rows):\n"
        text += df.head(50).to_string(index=False)
        return text
    except Exception as e:
        return f"Error reading Excel file: {str(e)}"

def extract_text_from_image(file):
    try:
        image = Image.open(file)
        text = f"Image File: {file.name}\n"
        text += f"Format: {image.format}\n"
        text += f"Size: {image.size[0]}x{image.size[1]} pixels\n"
        text += f"Mode: {image.mode}\n\n"
        text += "Note: Image uploaded. Ask questions about its content."
        return text
    except Exception as e:
        return f"Error reading image: {str(e)}"

def process_uploaded_file(uploaded_file):
    file_extension = uploaded_file.name.split('.')[-1].lower()
    if file_extension == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == 'docx':
        return extract_text_from_docx(uploaded_file)
    elif file_extension == 'txt':
        return extract_text_from_txt(uploaded_file)
    elif file_extension in ['xlsx', 'xls', 'csv']:
        return extract_text_from_excel(uploaded_file)
    elif file_extension in ['png', 'jpg', 'jpeg']:
        return extract_text_from_image(uploaded_file)
    else:
        return f"Unsupported file type: {file_extension}"

def get_azure_chatbot_response(messages, azure_endpoint, azure_key, deployment_name, api_version="2024-08-01-preview"):
    client = AzureOpenAI(
        azure_endpoint=azure_endpoint,
        api_key=azure_key,
        api_version=api_version
    )
    stream = client.chat.completions.create(
        model=deployment_name,
        messages=messages,
        stream=True,
        temperature=0.7,
    )
    return stream

# ------------------- MAIN APP -------------------
def main():
    st.set_page_config(page_title="TED Scraper & AI Assistant", layout="wide", initial_sidebar_state="collapsed")
    
    # Initialize language in session state
    if "language" not in st.session_state:
        st.session_state.language = "en"
    
    st.markdown("""
    <style>
    [data-testid="stAppViewContainer"] { background-color: #343541; }
    [data-testid="stHeader"] { background-color: #343541; }
    [data-testid="stSidebar"] { background-color: #202123; }
    .main .block-container { padding-top: 2rem !important; padding-bottom: 2rem !important; max-width: 48rem !important; margin: 0 auto !important; }
    .stChatMessage { background-color: transparent !important; padding: 1.5rem 0 !important; }
    [data-testid="stChatMessageContent"] { background-color: #444654 !important; border-radius: 0.5rem; padding: 1rem 1.5rem !important; color: #ececf1 !important; }
    [data-testid="stChatMessage"][data-testid*="user"] [data-testid="stChatMessageContent"] { background-color: #343541 !important; }
    .thinking-indicator { font-style: italic; color: #8e8ea0; font-size: 0.9rem; padding: 0.5rem 0; }
    .thinking-dots::after { content: '...'; animation: dots 1.5s steps(4, end) infinite; }
    @keyframes dots { 0%, 20% { content: '.'; } 40% { content: '..'; } 60%, 100% { content: '...'; } }
    [data-testid="stChatInput"] textarea { background-color: #40414f !important; color: #ececf1 !important; border: 1px solid #565869 !important; border-radius: 0.75rem !important; padding: 0.75rem 1rem !important; font-size: 1rem !important; min-height: 52px !important; max-height: 200px !important; line-height: 1.5 !important; resize: none !important; }
    [data-testid="stChatInput"] textarea:focus { border-color: #10a37f !important; box-shadow: 0 0 0 1px #10a37f !important; outline: none !important; }
    [data-testid="stChatInput"] > div { background-color: transparent !important; border: none !important; }
    .stMarkdown, .stText { color: #ececf1 !important; }
    h1, h2, h3, h4, h5, h6 { color: #ececf1 !important; }
    .stButton button { background-color: #10a37f !important; color: white !important; border: none !important; border-radius: 0.375rem !important; padding: 0.5rem 1rem !important; font-weight: 500 !important; }
    .stButton button:hover { background-color: #1a7f64 !important; }
    [data-testid="stChatMessage"] img { border-radius: 0.25rem !important; width: 32px !important; height: 32px !important; }
    .stSuccess, .stInfo, .stWarning { background-color: #444654 !important; color: #ececf1 !important; border-radius: 0.5rem !important; }
    .stTabs [data-baseweb="tab-list"] { gap: 2rem; background-color: #343541; border-bottom: 1px solid #565869; }
    .stTabs [data-baseweb="tab"] { color: #ececf1 !important; background-color: transparent; border-bottom: 2px solid transparent; padding: 1rem 0; font-weight: 500; }
    .stTabs [aria-selected="true"] { border-bottom-color: #10a37f !important; color: #10a37f !important; }
    .stTextInput input, .stDateInput input, .stSelectbox select, .stMultiSelect select { background-color: #40414f !important; color: #ececf1 !important; border: 1px solid #565869 !important; border-radius: 0.375rem !important; }
    label { color: #ececf1 !important; }
    .stDownloadButton button { background-color: #10a37f !important; color: white !important; }
    .streamlit-expanderHeader { background-color: #444654 !important; color: #ececf1 !important; border-radius: 0.5rem !important; }
    [data-testid="stFileUploader"] { background-color: transparent !important; border: none !important; padding: 0 !important; margin-bottom: 1rem !important; }
    [data-testid="stFileUploader"] section { border: 1px dashed #565869 !important; border-radius: 0.5rem !important; padding: 0.75rem !important; background-color: #40414f !important; }
    [data-testid="stFileUploader"] button { background-color: #565869 !important; color: #ececf1 !important; font-size: 0.875rem !important; padding: 0.25rem 0.5rem !important; }
    [data-testid="stDataFrame"] { background-color: #40414f !important; }
    </style>
    """, unsafe_allow_html=True)
    
    auth_flow()
    
    # Language switcher in top right corner
    header_col1, header_col2, header_col3 = st.columns([6, 1, 1])
    with header_col1:
        st.title(t("title"))
    with header_col3:
        selected_lang = st.selectbox(
            t("language"),
            options=["English", "Deutsch"],
            index=0 if st.session_state.language == "en" else 1,
            key="lang_selector",
            label_visibility="collapsed"
        )
        # Update language when changed
        new_lang = "en" if selected_lang == "English" else "de"
        if new_lang != st.session_state.language:
            st.session_state.language = new_lang
            st.rerun()
    
    if "scraped_data" not in st.session_state:
        st.session_state.scraped_data = None
    
    tab1, tab2 = st.tabs([t("tab_scraper"), t("tab_assistant")])
    
    # ============= TAB 1: TED SCRAPER =============
    with tab1:
        st.header(t("scraper_header"))
        st.write(t("scraper_subtitle"))
        
        with st.expander(t("instructions_header"), expanded=False):
            st.write(f"""
            {t("features_title")}
            {t("feature_1")}
            {t("feature_2")}
            {t("feature_3")}
            {t("feature_4")}
            
            {t("keywords_examples")}
            {t("example_1")}
            {t("example_2")}
            
            {t("multiselect_info")}
            {t("multiselect_1")}
            {t("multiselect_2")}
            """)

        st.subheader(t("search_criteria"))
        
        col1, col2 = st.columns(2)
        with col1:
            keywords = st.text_input(
                t("keywords_label"),
                placeholder=t("keywords_placeholder"),
                help=t("keywords_help")
            )
        with col2:
            cpv_codes = st.text_input(
                t("cpv_label"),
                "71541000 71500000 71240000 79421000 71000000 71248000 71312000 71700000 71300000 71520000 71250000 90712000 71313000",
                help=t("cpv_help")
            )

        col3, col4 = st.columns(2)
        with col3:
            buyer_country = st.text_input(t("country_label"), "DEU", help=t("country_help"))
        with col4:
            today = date.today()
            date_col1, date_col2 = st.columns(2)
            with date_col1:
                start_date_obj = st.date_input(t("date_start_label"), value=today)
            with date_col2:
                end_date_obj = st.date_input(t("date_end_label"), value=today)

        date_start = start_date_obj.strftime("%Y%m%d")
        date_end = end_date_obj.strftime("%Y%m%d")

        if st.button(t("search_button"), type="primary"):
            if not keywords.strip() and not cpv_codes.strip():
                st.error(t("error_no_keywords"))
            else:
                with st.spinner(t("searching")):
                    try:
                        rows = main_scraper(cpv_codes, keywords, date_start, date_end, buyer_country)
                        st.session_state.scraped_data = rows
                        if len(rows) > 0:
                            st.success(t("success_found", count=len(rows)))
                        else:
                            st.warning(t("warning_no_results"))
                    except Exception as e:
                        st.error(t("error_search", error=str(e)))
                        import traceback
                        st.code(traceback.format_exc())

        # Display results with MULTISELECT filtering
        if st.session_state.scraped_data:
            st.markdown("---")
            st.subheader(t("results_header"))
            
            df = pd.DataFrame(st.session_state.scraped_data)
            st.info(t("total_results", count=len(df)))
            
            with st.expander(t("filter_results"), expanded=True):
                filter_row1_col1, filter_row1_col2, filter_row1_col3 = st.columns(3)
                
                with filter_row1_col1:
                    if "Beschaffer" in df.columns:
                        beschaffer_options = sorted(df["Beschaffer"].dropna().unique().tolist())
                        selected_beschaffer = st.multiselect(
                            t("filter_beschaffer"),
                            options=beschaffer_options,
                            default=[],
                            help=t("filter_beschaffer_help")
                        )
                    else:
                        selected_beschaffer = []
                
                with filter_row1_col2:
                    if "Ort/Region" in df.columns:
                        region_options = sorted(df["Ort/Region"].dropna().unique().tolist())
                        selected_regions = st.multiselect(
                            t("filter_region"),
                            options=region_options,
                            default=[],
                            help=t("filter_region_help")
                        )
                    else:
                        selected_regions = []
                
                with filter_row1_col3:
                    if "Projektvolumen" in df.columns:
                        volume_filter = st.text_input(t("filter_volume"), placeholder=t("filter_volume_placeholder"))
                    else:
                        volume_filter = ""
                
                st.markdown(t("date_filters"))
                filter_row2_col1, filter_row2_col2, filter_row2_col3 = st.columns(3)
                
                with filter_row2_col1:
                    filter_projektstart = st.date_input(
                        t("filter_projektstart"),
                        value=None,
                        help=t("filter_projektstart_help")
                    )
                
                with filter_row2_col2:
                    filter_projektende = st.date_input(
                        t("filter_projektende"),
                        value=None,
                        help=t("filter_projektende_help")
                    )
                
                with filter_row2_col3:
                    filter_frist = st.date_input(
                        t("filter_frist"),
                        value=None,
                        help=t("filter_frist_help")
                    )
            
            # Apply filters
            filtered_df = df.copy()
            
            if selected_beschaffer:
                filtered_df = filtered_df[filtered_df["Beschaffer"].isin(selected_beschaffer)]
            
            if selected_regions:
                filtered_df = filtered_df[filtered_df["Ort/Region"].isin(selected_regions)]
            
            if volume_filter:
                try:
                    min_volume = float(volume_filter)
                    filtered_df["volume_numeric"] = filtered_df["Projektvolumen"].str.extract(r'([\d,.]+)')[0].str.replace(',', '').astype(float, errors='ignore')
                    filtered_df = filtered_df[filtered_df["volume_numeric"] >= min_volume]
                    filtered_df = filtered_df.drop(columns=["volume_numeric"])
                except:
                    st.warning(t("warning_volume"))
            
            if filter_projektstart:
                filtered_df["projektstart_date"] = pd.to_datetime(filtered_df["Projektstart"], errors='coerce')
                filtered_df = filtered_df[
                    (filtered_df["projektstart_date"].isna()) | 
                    (filtered_df["projektstart_date"] >= pd.Timestamp(filter_projektstart))
                ]
                filtered_df = filtered_df.drop(columns=["projektstart_date"])
            
            if filter_projektende:
                filtered_df["projektende_date"] = pd.to_datetime(filtered_df["Projektende"], errors='coerce')
                filtered_df = filtered_df[
                    (filtered_df["projektende_date"].isna()) | 
                    (filtered_df["projektende_date"] <= pd.Timestamp(filter_projektende))
                ]
                filtered_df = filtered_df.drop(columns=["projektende_date"])
            
            if filter_frist:
                filtered_df["frist_date"] = pd.to_datetime(filtered_df["Frist Abgabedatum"], errors='coerce')
                filtered_df = filtered_df[
                    (filtered_df["frist_date"].isna()) | 
                    (filtered_df["frist_date"] >= pd.Timestamp(filter_frist))
                ]
                filtered_df = filtered_df.drop(columns=["frist_date"])
            
            st.info(t("filtered_results", count=len(filtered_df)))
            
            st.dataframe(
                filtered_df,
                use_container_width=True,
                height=400,
                column_config={
                    "Ted-Link": st.column_config.LinkColumn("TED Link"),
                    "Vergabeplattform": st.column_config.LinkColumn("Platform")
                }
            )
            
            col_dl1, col_dl2 = st.columns(2)
            
            with col_dl1:
                if len(filtered_df) > 0:
                    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_excel:
                        try:
                            filtered_rows = filtered_df.to_dict('records')
                            save_to_excel(filtered_rows, temp_excel.name)
                            
                            with open(temp_excel.name, "rb") as f:
                                st.download_button(
                                    label=t("download_filtered", count=len(filtered_df)),
                                    data=f.read(),
                                    file_name=f"ted_filtered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    type="primary"
                                )
                        finally:
                            temp_excel.close()
                            if os.path.exists(temp_excel.name):
                                os.remove(temp_excel.name)
            
            with col_dl2:
                with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_excel:
                    try:
                        save_to_excel(st.session_state.scraped_data, temp_excel.name)
                        
                        with open(temp_excel.name, "rb") as f:
                            st.download_button(
                                label=t("download_all", count=len(df)),
                                data=f.read(),
                                file_name=f"ted_all_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                    finally:
                        temp_excel.close()
                        if os.path.exists(temp_excel.name):
                            os.remove(temp_excel.name)
    
    # ============= TAB 2: CHATBOT =============
    with tab2:
        with st.sidebar:
            azure_endpoint = get_secret("AZURE_ENDPOINT", "")
            azure_key = get_secret("AZURE_API_KEY", "")
            deployment_name = get_secret("DEPLOYMENT_NAME", "gpt-4o-mini")
            api_version = "2024-08-01-preview"
            
            st.markdown(t("config_header"))
            if azure_endpoint and azure_key:
                st.success(t("azure_connected"))
                try:
                    masked_endpoint = azure_endpoint.replace("https://", "").split(".")[0]
                    st.info(f"🔗 {masked_endpoint}")
                except:
                    st.info("🔗 Endpoint configured")
                st.info(f"🤖 {deployment_name}")
            else:
                st.warning(t("azure_warning"))
            
            st.markdown("---")
            st.markdown(t("doc_library"))
            st.caption(t("doc_optional"))
            
            if "document_store" not in st.session_state:
                st.session_state.document_store = {}
            
            library_files = st.file_uploader(
                "Upload Documents", 
                type=['pdf', 'docx', 'txt', 'xlsx', 'xls', 'csv', 'png', 'jpg', 'jpeg'],
                accept_multiple_files=True,
                key="library_uploader",
                help=t("file_help"),
                label_visibility="collapsed"
            )
            
            if library_files:
                for uploaded_file in library_files:
                    if uploaded_file.name not in st.session_state.document_store:
                        with st.spinner(t("processing", filename=uploaded_file.name)):
                            text = process_uploaded_file(uploaded_file)
                            if text:
                                st.session_state.document_store[uploaded_file.name] = text
                                st.success(t("file_added", filename=uploaded_file.name))
            
            if st.session_state.document_store:
                st.markdown(f"**📁 {len(st.session_state.document_store)} document(s)**")
                for doc_name in list(st.session_state.document_store.keys()):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.caption(f"• {doc_name}")
                    with col2:
                        if st.button("🗑️", key=f"del_{doc_name}"):
                            del st.session_state.document_store[doc_name]
                            st.rerun()
            
            st.markdown("---")
            if st.button(t("clear_chat"), use_container_width=True):
                st.session_state.chat_messages = []
                st.rerun()
        
        if not azure_endpoint or not azure_key:
            st.error(t("azure_error"))
            st.info(t("azure_info") + """
            ```
            AZURE_ENDPOINT = "https://your-resource.openai.azure.com"
            AZURE_API_KEY = "your-api-key"
            DEPLOYMENT_NAME = "gpt-4o-mini"
            ```
            """)
        else:
            if "chat_messages" not in st.session_state:
                st.session_state.chat_messages = []
            
            if not st.session_state.chat_messages:
                with st.chat_message("assistant", avatar=JKM_LOGO_URL):
                    st.markdown(f"""
                    {t("welcome_header")}
                    
                    {t("welcome_text")}
                    
                    {t("possibilities")}
                    {t("possibility_1")}
                    {t("possibility_2")}
                    {t("possibility_3")}
                    {t("possibility_4")}
                    
                    {t("ask_question")}
                    """)
            
            for message in st.session_state.chat_messages:
                avatar = JKM_LOGO_URL if message["role"] == "assistant" else BOT_AVATAR_URL
                with st.chat_message(message["role"], avatar=avatar):
                    st.markdown(message["content"])
            
            st.markdown("---")
            quick_file = st.file_uploader(
                t("file_upload"), 
                type=['pdf', 'docx', 'txt', 'xlsx', 'xls', 'csv', 'png', 'jpg', 'jpeg'],
                key="quick_uploader",
                help=t("file_help")
            )
            
            if quick_file:
                if quick_file.name not in st.session_state.document_store:
                    with st.spinner(t("processing", filename=quick_file.name)):
                        text = process_uploaded_file(quick_file)
                        if text:
                            st.session_state.document_store[quick_file.name] = text
                            st.success(t("file_added", filename=quick_file.name))
                            st.rerun()
            
            if prompt := st.chat_input(t("chat_input")):
                context_parts = []
                
                if st.session_state.document_store:
                    library_context = "\n\n".join([
                        f"=== DOCUMENT: {name} ===\n{content[:5000]}" 
                        for name, content in st.session_state.document_store.items()
                    ])
                    context_parts.append(library_context)
                
                st.session_state.chat_messages.append({"role": "user", "content": prompt})
                with st.chat_message("user", avatar=BOT_AVATAR_URL):
                    st.markdown(prompt)
                
                with st.chat_message("assistant", avatar=JKM_LOGO_URL):
                    thinking_placeholder = st.empty()
                    thinking_placeholder.markdown(f'<div class="thinking-indicator"><span class="thinking-dots">{t("thinking")}</span></div>', unsafe_allow_html=True)
                    
                    if context_parts:
                        full_context = "\n\n".join(context_parts)
                        system_content = f"""You are JKM AI Assistant - a helpful AI assistant for tenders, procurement documents, and general tasks.

You have access to the following documents:

{full_context}

INSTRUCTIONS:
- Analyze and answer questions based on the provided documents
- Extract specific information, identify empty fields, requirements, deadlines, etc.
- Always respond in German when asked in German, otherwise in English
- Be precise, professional, and helpful
- When analyzing PDFs: Look for specific sections, fields, tables, and requirements
- Summarize key information clearly"""
                    else:
                        system_content = """You are JKM AI Assistant - a helpful AI assistant for general questions and tasks.

INSTRUCTIONS:
- Answer general questions helpfully and precisely
- Always respond in German when asked in German, otherwise in English
- Be professional and friendly
- For procurement/tender questions: If documents are uploaded, analyze them in detail"""
                    
                    system_message = {"role": "system", "content": system_content}
                    
                    api_messages = [system_message] + [
                        {"role": m["role"], "content": m["content"]}
                        for m in st.session_state.chat_messages
                    ]
                    
                    try:
                        stream = get_azure_chatbot_response(
                            api_messages, 
                            azure_endpoint, 
                            azure_key, 
                            deployment_name,
                            api_version
                        )
                        
                        response_text = ""
                        for chunk in stream:
                            if hasattr(chunk, 'choices') and len(chunk.choices) > 0:
                                if hasattr(chunk.choices[0], 'delta') and hasattr(chunk.choices[0].delta, 'content'):
                                    if chunk.choices[0].delta.content:
                                        response_text += chunk.choices[0].delta.content
                        
                        thinking_placeholder.empty()
                        st.markdown(response_text)
                        
                        st.session_state.chat_messages.append({"role": "assistant", "content": response_text})
                        
                    except Exception as e:
                        thinking_placeholder.empty()
                        st.error(f"❌ Error: {str(e)}")
                        st.info(t("error_check_config"))

if __name__ == "__main__":
    main()
